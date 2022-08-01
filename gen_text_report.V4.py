import docx.text.paragraph
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime as dt
import os
import sys
import pandas as pd


def reformat_position_df(t_raw_df: pd.DataFrame) -> pd.DataFrame:
    t_raw_df = t_raw_df.fillna("--")
    t_raw_df["持仓数量"] = t_raw_df["持仓数量"].astype(int)
    t_raw_df["单位成本"] = t_raw_df["单位成本"].map(lambda z: z if z == "--" else "{:.2f}".format(z))
    t_raw_df["总成本"] = t_raw_df["总成本"].map(lambda z: "{:.2f}".format(z))
    t_raw_df["收盘价"] = t_raw_df["收盘价"].map(lambda z: z if z == "--" else "{:.2f}".format(z))
    t_raw_df["证券市值"] = t_raw_df["证券市值"].map(lambda z: "{:.2f}".format(z))
    t_raw_df["浮动盈亏"] = t_raw_df["浮动盈亏"].map(lambda z: "{:.2f}".format(z))
    t_raw_df["比例"] = t_raw_df["比例"].map(lambda z: z if z == "--" else "{:.2f}%".format(z * 100))
    return t_raw_df


def get_commodity_info(t_report_date: str, t_report_dir: str, t_header_row: int = 3) -> (str, pd.DataFrame):
    position_file = "04_衍生品持仓情况明细表_大宗商品_{}.xlsx".format(t_report_date)
    position_path = os.path.join(t_report_dir, report_date[0:4], report_date, position_file)
    position_df = pd.read_excel(position_path, dtype={"证券名称": str, "总成本": float}, header=t_header_row)
    position_df = position_df.dropna(axis=1, how="all")

    desc = ""
    for x in position_df["证券名称"].astype(str):
        if x.find("年初至今") > 0:
            desc = x.strip().replace(":", "：").replace(",", "，")
            desc = desc.replace("注：年初至今，", "相较去年末，")
            break
    print(desc)

    position_df = position_df.dropna(axis=0, how="any", subset=["持仓数量"])
    position_df = reformat_position_df(t_raw_df=position_df)
    return desc, position_df


def get_equity_desc_text(t_report_date: str, t_report_dir: str, t_header_row: int = 3) -> (str, pd.DataFrame):
    position_file = "04_持仓情况明细表_固收托管项目_{}.xlsx".format(t_report_date)
    position_path = os.path.join(t_report_dir, report_date[0:4], report_date, position_file)
    position_df = pd.read_excel(position_path, dtype={"证券名称": str, "总成本": float}, header=t_header_row)
    position_df = position_df.dropna(axis=1, how="all")

    cost_val_df = position_df.dropna(axis=0, how="all").set_index("证券名称")
    tot_cost_val = cost_val_df.at["合计", "总成本"]
    desc = "(1) 固收托管项目：今日无交易，持仓成本{:.2f}万元。".format(tot_cost_val / 1e4)
    print(desc)

    position_df = position_df.dropna(axis=0, how="any", subset=["持仓数量"])
    position_df = position_df.drop(axis=1, labels="账户")
    position_df = reformat_position_df(t_raw_df=position_df)
    return desc, position_df


def rgb_to_hex(t_color_code_rgb: tuple) -> str:
    hr = str(hex(int(t_color_code_rgb[0])))[-2:]
    hg = str(hex(int(t_color_code_rgb[1])))[-2:]
    hb = str(hex(int(t_color_code_rgb[2])))[-2:]
    return hr + hg + hb


# ================================= DOCX =================================
def indent_table(t_table: docx.oxml.table, t_indent: int):
    # noinspection PyProtectedMember
    tbl_pr = t_table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(t_indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)
    return 0


def set_table_cell_bg_color(t_table_cell: docx.oxml.table.CT_Tc, t_color_code_rgb: tuple):
    """
    set background shading for Header Rows
    """
    tbl_cell_properties = t_table_cell._element.tcPr
    cl_shading = OxmlElement('w:shd')
    cl_shading.set(qn('w:fill'), rgb_to_hex(t_color_code_rgb=t_color_code_rgb))  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tbl_cell_properties.append(cl_shading)
    return 0


def add_table_from_df(t_doc: Document, t_raw_df: pd.DataFrame, t_table_config: dict, t_table_indent: int = -1000, t_header_bg_color: tuple = (255, 255, 255)) -> int:
    nrow, ncol = t_raw_df.shape
    table = t_doc.add_table(rows=nrow + 1, cols=ncol, style="Table Grid")
    indent_table(table, t_table_indent)

    # update header
    hdr_cells = table.rows[0].cells
    for j in range(ncol):
        hdr_cells[j].text = t_raw_df.columns[j]
        hdr_cells[j].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        hdr_cells[j].width = Cm(t_table_config[t_raw_df.columns[j]]["width"])
        set_table_cell_bg_color(t_table_cell=hdr_cells[j], t_color_code_rgb=t_header_bg_color)

    # update contents
    for i in range(nrow):
        row_cells = table.rows[i + 1].cells
        for j in range(ncol):
            row_cells[j].text = str(t_raw_df.iloc[i, j])
            # set alignment
            if j > 0:
                row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            # set width
            if j == 0:
                row_cells[j].width = Cm(t_table_config[t_raw_df.columns[j]]["width"])

    # set table font
    table.style.font.name = "MSYH"
    table.style.font.size = Pt(8)
    table.style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # make table center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return 0


def change_run_fonts(t_run: docx.text.paragraph.Run, t_font_name: str, t_font_size: int, t_used_font: str):
    """

    :param t_run:
    :param t_font_name:
    :param t_font_size: positive integer, scale = Inches
    :param t_used_font:
    :return:
    """
    t_run.font.name = t_font_name
    t_run.font.size = Pt(t_font_size)
    t_run._element.rPr.rFonts.set(qn("w:eastAsia"), t_used_font)
    return 0


report_date = sys.argv[1]

# --- configure
commodity_src_dir = os.path.join("/Works", "Trade", "Reports", "output")
equity_src_dir = os.path.join("/Works", "Trade", "Reports_Merge", "output")
out_put_dir = os.path.join("C:", "Users", "Administrator", "Desktop")
# commodity_src_dir = "./data"
# equity_src_dir = "./data"
# out_put_dir = "."

table_config = {
    "证券名称": {"width": 3.0, },
    "证券代码": {"width": 1.6, },
    "持仓数量": {"width": 1.6, },
    "单位成本": {"width": 1.7, },
    "总成本": {"width": 2, },
    "收盘价": {"width": 1.7, },
    "证券市值": {"width": 2, },
    "浮动盈亏": {"width": 2, },
    "比例": {"width": 1.4, },
}
norm_font_size = 14
table_header_bg_color = (195, 195, 195)

# --- CREATE NEW
report = Document()
paragraph_count = -1

# # --- ADD HEAD
# report.add_heading("日报-{}".format(report_date), 1)
# paragraph_count += 1

# --- COMMODITY
commodity_desc_text, commodity_pos_df = get_commodity_info(t_report_date=report_date, t_report_dir=commodity_src_dir)
change_run_fonts(report.add_paragraph().add_run("3、衍生品大宗商品类："), "Fangsong", norm_font_size, "仿宋")
paragraph_count += 1
change_run_fonts(report.add_paragraph().add_run(commodity_desc_text), "Fangsong", norm_font_size, "仿宋")
paragraph_count += 1
report.paragraphs[paragraph_count].paragraph_format.first_line_indent = Pt(norm_font_size * 2)
add_table_from_df(t_doc=report, t_raw_df=commodity_pos_df, t_table_config=table_config, t_header_bg_color=table_header_bg_color)

report.add_page_break()

# --- EQUITY
equity_desc_text, equity_pos_df = get_equity_desc_text(t_report_date=report_date, t_report_dir=equity_src_dir)
change_run_fonts(report.add_paragraph().add_run("5、协作类："), "Fangsong", norm_font_size, "仿宋")
paragraph_count += 1
change_run_fonts(report.add_paragraph().add_run(equity_desc_text), "Fangsong", norm_font_size, "仿宋")
paragraph_count += 1
report.paragraphs[paragraph_count].paragraph_format.first_line_indent = Pt(norm_font_size * 2)
add_table_from_df(t_doc=report, t_raw_df=equity_pos_df, t_table_config=table_config, t_header_bg_color=table_header_bg_color)

# --- SAVE
report_file = "日报_{}.docx".format(report_date)
report_path = os.path.join(out_put_dir, report_file)
report.save(report_path)
print("... {} {} created.".format(dt.datetime.now(), report_path))
